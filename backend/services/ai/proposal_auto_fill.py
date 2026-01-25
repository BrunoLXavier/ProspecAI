# Implements RF-08: Auto-fill proposal fields from uploaded documents via AI
"""
Proposal Auto-Fill Service

This service handles:
1. Text extraction from PDF/DOCX files
2. LLM-based field mapping
3. Kafka messaging for async processing
4. WebSocket notifications for real-time updates

Uses pdfplumber (PDF) and python-docx (DOCX) for extraction.
Uses configured LLM provider for field mapping.
"""
from typing import Dict, Any, List, Optional
from uuid import UUID
from datetime import datetime
import logging
import json
import asyncio

logger = logging.getLogger(__name__)

# Lazy imports for document processing (may not be available in all environments)
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber not available - PDF extraction disabled")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX extraction disabled")


class ProposalAutoFillService:
    """
    Service for auto-filling proposal fields from uploaded documents.
    Implements RF-08: AI-assisted data extraction with human-in-the-loop.
    """
    
    KAFKA_TOPIC_REQUEST = "prospecai.proposal.auto-fill-request"
    KAFKA_TOPIC_RESULTS = "prospecai.proposal.auto-fill-results"
    
    def __init__(
        self,
        kafka_producer=None,
        llm_service=None,
        file_storage=None,
        websocket_manager=None,
        attachment_repository=None,
        suggestion_repository=None,
        field_template_repository=None
    ):
        self.kafka_producer = kafka_producer
        self.llm_service = llm_service
        self.file_storage = file_storage
        self.websocket_manager = websocket_manager
        self.attachment_repository = attachment_repository
        self.suggestion_repository = suggestion_repository
        self.field_template_repository = field_template_repository
    
    async def request_extraction(
        self,
        attachment_id: UUID,
        proposal_id: UUID,
        file_key: str,
        file_type: str,
        tenant_id: UUID
    ) -> None:
        """
        Publish extraction request to Kafka for async processing.
        """
        message = {
            "attachment_id": str(attachment_id),
            "proposal_id": str(proposal_id),
            "file_key": file_key,
            "file_type": file_type,
            "tenant_id": str(tenant_id),
            "requested_at": datetime.utcnow().isoformat()
        }
        
        if self.kafka_producer:
            await self.kafka_producer.send_and_wait(
                self.KAFKA_TOPIC_REQUEST,
                json.dumps(message).encode("utf-8")
            )
            logger.info(f"Published auto-fill request for attachment {attachment_id}")
        else:
            # Fallback: process synchronously
            logger.warning("Kafka not available, processing synchronously")
            await self.process_extraction(message)
    
    async def process_extraction(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process extraction request (called by Kafka consumer or directly).
        """
        attachment_id = UUID(message["attachment_id"])
        proposal_id = UUID(message["proposal_id"])
        file_key = message["file_key"]
        file_type = message["file_type"]
        tenant_id = UUID(message["tenant_id"])
        
        logger.info(f"Processing auto-fill for attachment {attachment_id}")
        
        try:
            # Update attachment status
            if self.attachment_repository:
                attachment = await self.attachment_repository.get_by_id(attachment_id, tenant_id)
                if attachment:
                    attachment.start_extraction()
                    await self.attachment_repository.update(attachment)
            
            # Download file from MinIO
            file_content = await self.file_storage.download(file_key)
            
            # Extract text based on file type
            extracted_text = await self._extract_text(file_content, file_type)
            
            if not extracted_text:
                raise ValueError(f"Could not extract text from file type: {file_type}")
            
            # Get field templates for this proposal
            field_templates = await self._get_proposal_field_templates(proposal_id, tenant_id)
            
            # Use LLM to map extracted text to fields
            field_suggestions = await self._map_text_to_fields(
                extracted_text, field_templates, tenant_id
            )
            
            # Save suggestions
            saved_suggestions = []
            for suggestion in field_suggestions:
                saved = await self._save_suggestion(
                    proposal_id=proposal_id,
                    attachment_id=attachment_id,
                    field_key=suggestion["field_key"],
                    suggested_value=suggestion["value"],
                    confidence_score=suggestion["confidence"],
                    source_text=suggestion.get("source_text"),
                    source_page=suggestion.get("source_page"),
                    tenant_id=tenant_id
                )
                saved_suggestions.append(saved)
            
            # Update attachment with results
            if self.attachment_repository:
                attachment = await self.attachment_repository.get_by_id(attachment_id, tenant_id)
                if attachment:
                    attachment.complete_extraction(
                        text=extracted_text[:10000],  # Store first 10k chars
                        fields={s["field_key"]: s["value"] for s in field_suggestions}
                    )
                    await self.attachment_repository.update(attachment)
            
            # Notify via WebSocket
            await self._notify_extraction_complete(
                proposal_id=proposal_id,
                attachment_id=attachment_id,
                suggestions_count=len(saved_suggestions),
                tenant_id=tenant_id
            )
            
            logger.info(f"Extraction complete for {attachment_id}: {len(saved_suggestions)} suggestions")
            
            return {
                "status": "success",
                "suggestions_count": len(saved_suggestions),
                "attachment_id": str(attachment_id)
            }
            
        except Exception as e:
            logger.error(f"Extraction failed for {attachment_id}: {e}")
            
            # Update attachment with error
            if self.attachment_repository:
                attachment = await self.attachment_repository.get_by_id(attachment_id, tenant_id)
                if attachment:
                    attachment.fail_extraction(str(e))
                    await self.attachment_repository.update(attachment)
            
            # Notify via WebSocket
            await self._notify_extraction_failed(
                proposal_id=proposal_id,
                attachment_id=attachment_id,
                error=str(e),
                tenant_id=tenant_id
            )
            
            return {
                "status": "failed",
                "error": str(e),
                "attachment_id": str(attachment_id)
            }
    
    async def _extract_text(self, file_content: bytes, file_type: str) -> Optional[str]:
        """
        Extract text from file content based on file type.
        """
        file_type_lower = file_type.lower()
        
        if "pdf" in file_type_lower:
            return await self._extract_pdf_text(file_content)
        elif "docx" in file_type_lower or "word" in file_type_lower:
            return await self._extract_docx_text(file_content)
        elif "text" in file_type_lower:
            return file_content.decode("utf-8", errors="ignore")
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return None
    
    async def _extract_pdf_text(self, file_content: bytes) -> Optional[str]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            logger.error("pdfplumber not available")
            return None
        
        try:
            import io
            text_parts = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return None
    
    async def _extract_docx_text(self, file_content: bytes) -> Optional[str]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return None
        
        try:
            import io
            doc = DocxDocument(io.BytesIO(file_content))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None
    
    async def _get_proposal_field_templates(
        self, proposal_id: UUID, tenant_id: UUID
    ) -> List[Dict[str, Any]]:
        """Get field templates for this proposal."""
        # For now, return standard fields + any template-specific
        from domain.entities.proposal import STANDARD_PROPOSAL_FIELDS
        
        fields = []
        for field in STANDARD_PROPOSAL_FIELDS:
            fields.append({
                "field_key": field["field_key"],
                "label": field["label"],
                "field_type": field["field_type"].value if hasattr(field["field_type"], "value") else field["field_type"],
                "help_text": field.get("help_text", ""),
                "auto_fill_prompt": field.get("auto_fill_prompt")
            })
        
        return fields
    
    async def _map_text_to_fields(
        self,
        text: str,
        field_templates: List[Dict[str, Any]],
        tenant_id: UUID
    ) -> List[Dict[str, Any]]:
        """
        Use LLM to extract field values from text.
        """
        if not self.llm_service:
            logger.warning("LLM service not available, skipping field mapping")
            return []
        
        # Build prompt with field definitions
        fields_description = "\n".join([
            f"- {f['field_key']}: {f['label']} ({f['field_type']})"
            + (f" - {f['help_text']}" if f.get("help_text") else "")
            for f in field_templates
        ])
        
        prompt = f"""Você é um assistente especializado em análise de documentos de propostas de P&D.

Analise o texto abaixo e extraia informações para preencher os campos de uma proposta.

CAMPOS DISPONÍVEIS:
{fields_description}

TEXTO DO DOCUMENTO:
{text[:8000]}

Para cada campo que você conseguir identificar no texto, retorne um JSON com:
- field_key: chave do campo
- value: valor extraído
- confidence: confiança de 0.0 a 1.0
- source_text: trecho do texto de onde extraiu (máximo 200 caracteres)

Retorne APENAS um array JSON válido, sem explicações. Exemplo:
[{{"field_key": "title", "value": "Título do Projeto", "confidence": 0.95, "source_text": "..."}}]

Se não encontrar nenhum campo, retorne: []
"""
        
        try:
            response = await self.llm_service.generate(
                prompt=prompt,
                tenant_id=tenant_id,
                max_tokens=2000,
                temperature=0.1
            )
            
            # Parse JSON response
            response_text = response.get("content", "[]")
            
            # Extract JSON from response (handle markdown code blocks)
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            suggestions = json.loads(response_text.strip())
            
            # Validate and normalize
            valid_suggestions = []
            for s in suggestions:
                if isinstance(s, dict) and "field_key" in s and "value" in s:
                    valid_suggestions.append({
                        "field_key": s["field_key"],
                        "value": s["value"],
                        "confidence": min(max(float(s.get("confidence", 0.5)), 0.0), 1.0),
                        "source_text": str(s.get("source_text", ""))[:200]
                    })
            
            return valid_suggestions
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM response as JSON: {e}")
            return []
        except Exception as e:
            logger.error(f"LLM field mapping failed: {e}")
            return []
    
    async def _save_suggestion(
        self,
        proposal_id: UUID,
        attachment_id: UUID,
        field_key: str,
        suggested_value: Any,
        confidence_score: float,
        source_text: Optional[str],
        source_page: Optional[int],
        tenant_id: UUID
    ) -> Dict[str, Any]:
        """Save auto-fill suggestion to database."""
        from domain.entities.proposal import ProposalFieldValue
        from uuid import uuid4
        
        suggestion_data = {
            "id": uuid4(),
            "proposal_id": proposal_id,
            "attachment_id": attachment_id,
            "field_key": field_key,
            "suggested_value": suggested_value,
            "confidence_score": confidence_score,
            "source_text": source_text,
            "source_page": source_page,
            "status": "pending",
            "tenant_id": tenant_id
        }
        
        if self.suggestion_repository:
            # Save to database (implement based on repository pattern)
            pass
        
        return suggestion_data
    
    async def _notify_extraction_complete(
        self,
        proposal_id: UUID,
        attachment_id: UUID,
        suggestions_count: int,
        tenant_id: UUID
    ) -> None:
        """Send WebSocket notification about extraction completion."""
        if self.websocket_manager:
            await self.websocket_manager.broadcast_to_room(
                room=f"proposal:{proposal_id}",
                message={
                    "event": "auto_fill_complete",
                    "proposal_id": str(proposal_id),
                    "attachment_id": str(attachment_id),
                    "suggestions_count": suggestions_count,
                    "timestamp": datetime.utcnow().isoformat()
                }
            )
    
    async def _notify_extraction_failed(
        self,
        proposal_id: UUID,
        attachment_id: UUID,
        error: str,
        tenant_id: UUID
    ) -> None:
        """Send WebSocket notification about extraction failure."""
        if self.websocket_manager:
            await self.websocket_manager.broadcast_to_room(
                room=f"proposal:{proposal_id}",
                message={
                    "event": "auto_fill_failed",
                    "proposal_id": str(proposal_id),
                    "attachment_id": str(attachment_id),
                    "error": error,
                    "timestamp": datetime.utcnow().isoformat()
                }
            )


class AutoFillKafkaConsumer:
    """
    Kafka consumer for processing auto-fill requests.
    Runs as background task or separate process.
    """
    
    def __init__(self, auto_fill_service: ProposalAutoFillService, kafka_consumer=None):
        self.auto_fill_service = auto_fill_service
        self.kafka_consumer = kafka_consumer
        self._running = False
    
    async def start(self):
        """Start consuming messages."""
        if not self.kafka_consumer:
            logger.warning("Kafka consumer not configured")
            return
        
        self._running = True
        logger.info("Starting auto-fill Kafka consumer")
        
        async for msg in self.kafka_consumer:
            if not self._running:
                break
            
            try:
                message = json.loads(msg.value.decode("utf-8"))
                await self.auto_fill_service.process_extraction(message)
            except Exception as e:
                logger.error(f"Failed to process auto-fill message: {e}")
    
    async def stop(self):
        """Stop consuming messages."""
        self._running = False
        if self.kafka_consumer:
            await self.kafka_consumer.stop()
